import os
import time
from pathlib import Path
from docx import Document
import torch
from shutil import copy2, rmtree
from TTS.api import TTS
from TTS.utils.synthesizer import Synthesizer
import io
from contextlib import redirect_stdout
import streamlit as st
import psutil
import re

# =============================
# CONSOLE STATUS / LOG HELPERS
# =============================
_last_standby_print = 0.0

def _console(status: str, msg: str = ""):
    """Simple console logger with timestamp and status flag."""
    ts = time.strftime('%Y-%m-%d %H:%M:%S')
    line = f"[{ts}] [{status.upper()}] {msg}"
    try:
        print(line, flush=True)
    except Exception:
        pass

def _console_standby_throttled(msg: str, interval: float = 8.0):
    """Avoid spamming STANDBY logs on Streamlit reruns."""
    global _last_standby_print
    now = time.time()
    if now - _last_standby_print >= interval:
        _console("STANDBY", msg)
        _last_standby_print = now

# =============================
# CONFIGURAÇÕES GERAIS
# =============================
# Prioridade de processo (Windows)
p = psutil.Process(os.getpid())
try:
    p.nice(psutil.HIGH_PRIORITY_CLASS)  # ou psutil.REALTIME_PRIORITY_CLASS
except Exception:
    pass

os.environ["COQUI_TOS_AGREED"] = "1"

DEFAULT_MODEL = "tts_models/multilingual/multi-dataset/xtts_v2"
DEFAULT_DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

# Caminho padrão para salvar os áudios gerados
DOWNLOAD_PATH = r"C:\\Users\\dud\\Downloads\\youtube\\VIAJENS\\AUDIOS_FEITOS"

# Lista para armazenar os textos problemáticos (quando há warning de limite)
textos_com_erro = []

# =============================
# UTILITÁRIOS
# =============================

def sanitize_name(name: str) -> str:
    """Remove caracteres problemáticos de nomes de arquivo/pasta."""
    base = re.sub(r"[\\/:*?\"<>|]", "_", name)
    base = re.sub(r"\s+", " ", base).strip()
    return base


def list_wav_files(folder: str):
    try:
        p = Path(folder)
        if not p.exists():
            return []
        return [str(f.name) for f in p.glob("*.wav")]
    except Exception:
        return []


# Substitui a função de segmentação para remover ponto final simples
params = {
    "remove_trailing_dots": True,
    "voice": "",  # definido via UI
    "language": "pt",  # definido via UI
    "model_name": DEFAULT_MODEL,
    "device": DEFAULT_DEVICE,
}


def new_split_into_sentences(self, text):
    sentences = self.seg.segment(text)
    if params['remove_trailing_dots']:
        sentences_without_dots = []
        for sentence in sentences:
            if sentence.endswith('.') and not sentence.endswith('...'):
                sentence = sentence[:-1]
            sentences_without_dots.append(sentence)
        return sentences_without_dots
    else:
        return sentences

# Aplica override na lib
Synthesizer.split_into_sentences = new_split_into_sentences


# =============================
# CARREGAMENTO DO MODELO
# =============================
@st.cache_resource(show_spinner=True)
def load_model(model_name: str, device: str):
    return TTS(model_name).to(device)


# =============================
# LEITURA DO DOCX
# =============================
@st.cache_data(show_spinner=False)
def load_text(file_path):
    doc = Document(file_path)
    return [paragraph.text.replace('.', ',.') for paragraph in doc.paragraphs if paragraph.text.strip()]


def generate_audio_filename(index):
    return f"audio_{index + 1}.wav"


# =============================
# GERAÇÃO COM LOG
# =============================

def tts_to_file_logged(model, text: str, out_path: str, language: str, speaker_wav: str, speed: float = 0.85):
    output_buffer = io.StringIO()
    with redirect_stdout(output_buffer):
        model.tts_to_file(
            text=text,
            file_path=out_path,
            speaker_wav=speaker_wav,
            language=language,
            speed=speed,
        )
    return output_buffer.getvalue()


# =============================
# APP STREAMLIT
# =============================

def main_app():
    _console_standby_throttled("Aplicação iniciada ouciando eventos da UI...")
    st.set_page_config(page_title="Gerador de Áudio de Viagem (WEB)", layout="wide")

    st.title("Gerador de Áudio de Viagem (WEB)")
    _console_standby_throttled("Interface carregada; aguardando ação do usuário.")
    st.caption("Agora com seleção de idioma (PT/ES), escolha de voz por pasta e preview instantâneo.")

    with st.sidebar:
        st.subheader("Configurações")

        # Idioma: PT ou ES
        language = st.radio("Idioma da narração", options=["pt", "es"], index=0, horizontal=True)
        params["language"] = language

        # Caminho da pasta de vozes
        default_voice_dir = r"C:\\Users\\dud\\Downloads\\youtube\\scripts\\voices"
        voice_dir = st.text_input("Pasta com vozes (.wav)", value=default_voice_dir)
        wavs = list_wav_files(voice_dir)
        if not wavs:
            st.warning("Nenhum arquivo .wav encontrado nessa pasta.")
        else:
            selected_wav_name = st.selectbox("Escolha a voz (arquivo .wav)", wavs)
            selected_wav_path = str(Path(voice_dir) / selected_wav_name)
            params["voice"] = selected_wav_path

            # Preview da voz: toca o próprio arquivo .wav
            st.markdown("**Pré-escuta da voz selecionada**")
            try:
                with open(selected_wav_path, "rb") as f:
                    st.audio(f.read(), format="audio/wav")
            except Exception as e:
                st.error(f"Erro ao carregar preview da voz: {e}")

        # Modelo e dispositivo (opcionais)
        with st.expander("Opções avançadas"):
            params["model_name"] = st.text_input("Modelo Coqui", value=params["model_name"])  # normalmente xtts_v2
            params["device"] = st.selectbox("Dispositivo", ["cuda", "cpu"], index=0 if DEFAULT_DEVICE == "cuda" else 1)
            remove_trailing = st.checkbox("Remover ponto final simples de frases", value=True)
            params["remove_trailing_dots"] = remove_trailing

    # Upload do DOCX
    st.write("Faça upload do arquivo .docx para gerar os áudios com logs ao vivo.")
    uploaded_file = st.file_uploader("Selecione o arquivo .docx", type="docx")

    # Preview rápido de amostra de TTS (sem precisar do DOCX)
    st.divider()
    st.markdown("### Teste rápido da voz selecionada (TTS)")
    sample_text = st.text_input("Texto de teste", value="Este é um teste de narração.")
    if st.button("Reproduzir amostra gerada"):
        _console("EXECUTANDO", "Gerando amostra TTS com a voz selecionada...")
        if not params.get("voice"):
            st.error("Selecione uma voz primeiro.")
        else:
            with st.spinner("Gerando amostra..."):
                model = load_model(params["model_name"], params["device"])  # cache_resource evita recarregar
                tmp_sample = Path("./_tmp_sample.wav")
                try:
                    _ = tts_to_file_logged(
                        model,
                        text=sample_text,
                        out_path=str(tmp_sample),
                        language=params["language"],
                        speaker_wav=params["voice"],
                        speed=0.9,
                    )
                    _console("EXECUTANDO", "Amostra gerada com sucesso, reproduzindo...")
                    with open(tmp_sample, "rb") as f:
                        st.audio(f.read(), format="audio/wav")
                finally:
                    if tmp_sample.exists():
                        tmp_sample.unlink(missing_ok=True)

    st.divider()

    if uploaded_file is not None:
        _console("EXECUTANDO", f"Arquivo recebido: {uploaded_file.name}. Preparando pastas e carregando parágrafos...")
        # Usa o nome real do arquivo enviado (sem "temp...") para a pasta de saída
        original_name = sanitize_name(uploaded_file.name)
        base_name = os.path.splitext(original_name)[0]
        lang_prefix = params["language"]  # 'pt' ou 'es'
        new_folder_name = f"{lang_prefix}_{base_name}"
        new_folder_path = os.path.join(DOWNLOAD_PATH, new_folder_name)
        os.makedirs(new_folder_path, exist_ok=True)

        # Salva o DOCX original dentro da pasta criada, mantendo o nome original
        docx_destination = os.path.join(new_folder_path, original_name)
        with open(docx_destination, "wb") as f:
            f.write(uploaded_file.read())

        st.success(f"Pasta criada: {new_folder_path}")

        # Carrega parágrafos
        paragraphs = load_text(docx_destination)

        # Carrega modelo (cacheado)
        model = load_model(params["model_name"], params["device"])  # evita recarga nas próximas execuções
        _console("EXECUTANDO", f"Modelo carregado ({params['model_name']} em {params['device']}). Iniciando geração de áudios...")


        st.write("Iniciando geração dos áudios...")
        progress = st.progress(0)
        log_placeholder = st.empty()
        textos_com_erro.clear()
        error_texts = []

        for index, paragraph in enumerate(paragraphs):
            output_file_name = generate_audio_filename(index)
            output_file_path = os.path.join(new_folder_path, output_file_name)

            st.write(f"Gerando áudio para o parágrafo {index + 1}/{len(paragraphs)}...")
            _console("EXECUTANDO", f"Parágrafo {index + 1}/{len(paragraphs)}: gerando arquivo {output_file_name}...")

            try:
                log = tts_to_file_logged(
                    model,
                    text=paragraph,
                    out_path=output_file_path,
                    language=params["language"],
                    speaker_wav=params["voice"],
                    speed=0.85,
                )
                log_placeholder.code(log or "(sem logs)")
                _console("EXECUTANDO", f"Parágrafo {index + 1}/{len(paragraphs)} finalizado: {output_file_name}")
            except Exception as e:
                log_placeholder.error(f"Erro ao gerar áudio para {output_file_name}: {e}")
                _console("ERRO", f"Falha no parágrafo {index + 1}/{len(paragraphs)} ({output_file_name}): {e}")
                if "exceeds the character limit" in str(e).lower():
                    error_texts.append(paragraph)
                continue

            # Checa warning de limite no stdout
            if "exceeds the character limit" in (log or "").lower():
                novo_nome = f"audio_{index + 1}__pode ter erro.wav"
                novo_caminho = os.path.join(new_folder_path, novo_nome)
                try:
                    os.rename(output_file_path, novo_caminho)
                except Exception:
                    pass
                textos_com_erro.append(paragraph)

            progress.progress((index + 1) / max(1, len(paragraphs)))

        # Se houve textos problemáticos, salva um docx de relatório
        if error_texts or textos_com_erro:
            error_docx_path = os.path.join(new_folder_path, "paragrafos_com_erro.docx")
            error_doc = Document()
            for t in error_texts + textos_com_erro:
                error_doc.add_paragraph(t)
            error_doc.save(error_docx_path)
            st.warning(f"Parágrafos com possível erro salvos em: {error_docx_path}")

        st.success("Processo de geração de áudio concluído!")
        _console("CONCLUÍDO", "Geração finalizada. Sistema em standby aguardando nova ação.")

    # Se nenhum arquivo foi enviado nesta execução do app, loga STANDBY no console (com throttling)
    if uploaded_file is None:
        _console_standby_throttled("Sem tarefas ativas. Aguardando upload de DOCX, mudança de opções ou clique em botões.")


if __name__ == "__main__":
    main_app()
